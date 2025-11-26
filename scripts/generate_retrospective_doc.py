#!/usr/bin/env python3
"""
Generate a formatted Word document for the Sprint 3 Review & Retrospective.

This script reads the Markdown source from `retrospective_assignment.md`
and writes a styled `.docx` document using python-docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = REPO_ROOT / "retrospective_assignment.md"
OUTPUT_PATH = REPO_ROOT / "Sprint3_Review_Retrospective.docx"


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    run.font.size = Pt(16 if level == 1 else 14)
    if level == 1:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix:
        run = paragraph.add_run(f"{bold_prefix} ")
        run.bold = True
        run.font.size = Pt(12)
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.size = Pt(12)


def main() -> None:
    source_text = MARKDOWN_PATH.read_text(encoding="utf-8")
    document = Document()

    add_heading(document, "Sprint 3 Review & Retrospective", level=1)

    document.add_paragraph()
    add_heading(document, "Sprint Review", level=2)

    add_paragraph(
        document,
        "Deliver a demo-ready SpendWise dashboard with actionable analytics, complete the "
        "Savings Circles feature set, and support the implementation with updated design documentation "
        "and unit tests.",
        bold_prefix="Sprint Goal.",
    )

    add_paragraph(
        document,
        "We committed in the pre-sprint agreement to: (1) embed MPAndroidChart visualizations showing "
        "category spend and budget usage, (2) finish the Savings Circles invitation and progress workflow "
        "backed by Firestore, (3) refresh the design packet (DCD, SD, SOLID, pattern evidence), and "
        "(4) add unit coverage for analytics aggregation and group rollover logic. All four objectives "
        "shipped. The dashboard now renders both pie and bar charts sourced from LiveData exposed by "
        "`DashboardAnalyticsViewModel`. The Savings Circles flow supports sending, accepting, and "
        "declining invitations, and the `SavingCircleViewModel` keeps member balances and cycle history "
        "synchronized. Design artifacts were merged into the shared drive and linked in the sprint "
        "deliverable. Two new unit test suites (`AnalyticsRepositoryTest`, `MemberCycleTest`) run green "
        "locally, covering data aggregation and rollover math.",
        bold_prefix="Objectives vs. Outcomes.",
    )

    add_paragraph(
        document,
        "Work stayed aligned with the ownership we established. Analytics tasks (chart integration, "
        "seeded fallbacks, view model plumbing) were handled by the dashboard sub-team; Savings Circles "
        "(invitation storage, lifecycle cleanup, UI polish) stayed with the social features pair; testing "
        "and design documentation were split across the remaining members. The only deviation was "
        "shifting UI polish for invitation empty-states to mid-sprint after Firestore schema tweaks "
        "delayed backend work; we agreed on the change during the Wednesday stand-up and absorbed the "
        "delay without impacting the demo scope.",
        bold_prefix="Adherence to Pre-Sprint Agreement.",
    )

    add_paragraph(
        document,
        "The dashboard now provides real-time visibility into spending habits, enabling users to adjust "
        "budgets proactively. Savings Circles unlock collaborative savings challenges with clear progress "
        "tracking, differentiating SpendWise from standard budgeting apps. Updated design artifacts keep "
        "the documentation synchronized with the codebase, and the new unit tests give us automated "
        "regression checks around the most logic-heavy features.",
        bold_prefix="Completed Deliverables & Value.",
    )

    document.add_paragraph()
    add_heading(document, "Sprint Retrospective", level=2)

    add_paragraph(
        document,
        "Communication was generally strong—Slack updates remained steady, and daily stand-ups surfaced "
        "blockers quickly. Pair programming between the analytics and repository owners shortened "
        "integration time. We did encounter some friction around coordinating Firestore rules changes; "
        "clarifying ownership during mid-sprint planning mitigated confusion.",
        bold_prefix="Team Dynamics.",
    )

    add_paragraph(
        document,
        "Early alignment on data schemas let us stub repositories and unblock UI work. The strategy/factory "
        "abstractions around chart configuration made it easy to experiment with additional visualizations. "
        "Sharing test utilities increased coverage without duplicating fixtures. Hosting mid-sprint demos "
        "fostered shared understanding of progress.",
        bold_prefix="What Worked Well.",
    )

    add_paragraph(
        document,
        "We underestimated the time required to refactor existing ViewModels to reuse analytics helpers, "
        "leading to a crunch late in the sprint. Design artifact updates were started near the deadline, "
        "creating parallel work that could have been spread earlier. Finally, onboarding new contributors "
        "to Firestore emulators needs a clearer guide; local setup problems slowed testing for a day.",
        bold_prefix="Areas for Improvement.",
    )

    add_paragraph(document, "", bold_prefix="Actionable Steps.")
    add_bullet(
        document,
        "Front-load design/documentation tasks by scheduling a dedicated working session in the first half "
        "of Sprint 4.",
    )
    add_bullet(
        document,
        "Extend the onboarding README with explicit Firestore emulator setup instructions and sample data "
        "seeding scripts.",
    )
    add_bullet(
        document,
        "Add setup scripts to automate provisioning of seeded analytics data so new test cases can be "
        "written quickly.",
    )
    add_bullet(
        document,
        "Reserve a mid-sprint hour for refactoring debt, ensuring shared utilities are stabilized before UI "
        "integration picks them up.",
    )

    add_paragraph(
        document,
        "The team delivered the agreed scope, tightened collaboration across feature boundaries, and "
        "produced customer-visible improvements. With earlier documentation updates and better tooling "
        "around environment setup, we expect to move faster in Sprint 4.",
        bold_prefix="Overall Sentiment.",
    )

    document.save(OUTPUT_PATH)
    print(f"Generated {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

